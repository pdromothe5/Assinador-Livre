import tkinter as tk
from tkinter import filedialog, messagebox
import PyPDF2
import hashlib
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import io
from datetime import datetime
from PIL import Image
from docx import Document
from tkinter.ttk import Progressbar


class PDFSignerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Assinador Livre")

        self.file_path = ""
        self.auth_code = ""
        self.image_file_path = None

        self.signed_hashes = []

        self.file_label = tk.Label(root, text="Selecione o arquivo (PDF or DOCX):")
        self.file_label.pack()

        self.file_button = tk.Button(root, text="Browser", command=self.browse_file)
        self.file_button.pack()

        self.signer_name_label = tk.Label(root, text="Insira o seu nome:")
        self.signer_name_label.pack()

        self.signer_name_entry = tk.Entry(root)
        self.signer_name_entry.pack()

        self.generate_hash_button = tk.Button(root, text="Gerar Hash", command=self.generate_hash)
        self.generate_hash_button.pack()

        self.sign_file_button = tk.Button(root, text="Assinar documento", command=self.sign_file)
        self.sign_file_button.pack()

        self.authenticator_label = tk.Label(root, text="Insira a hash para autenticar:")
        self.authenticator_label.pack()

        self.authenticator_entry = tk.Entry(root)
        self.authenticator_entry.pack()

        self.authenticate_button = tk.Button(root, text="Autenticar", command=self.authenticate_hash)
        self.authenticate_button.pack()

        self.image_label = tk.Label(root, text="Select a transparent image for signature background:")
        self.image_label.pack()

        self.image_button = tk.Button(root, text="Browser", command=self.browse_image)
        self.image_button.pack()

        self.result_label = tk.Label(root, text="")
        self.result_label.pack()

        self.clear_button = tk.Button(root, text="Limpar", command=self.clear_fields)
        self.clear_button.pack()
        
        self.progress_bar = Progressbar(root, orient="horizontal", length=300, mode="determinate")
        self.progress_bar.pack()

    
        self.root.mainloop()

    def clear_fields(self):
        self.file_path = ""
        self.auth_code = ""
        self.image_file_path = None
        self.signer_name_entry.delete(0, "end")
        self.authenticator_entry.delete(0, "end")
        self.result_label.config(text="")
        self.progress_bar["value"] = 0
        self.root.update()


    def browse_file(self):
        self.file_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf"), ("DOCX Files", "*.docx")])

    def generate_hash(self):
        signer_name = self.signer_name_entry.get()
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        data_to_hash = f"{signer_name}{timestamp}"

        hash_object = hashlib.sha256(data_to_hash.encode())
        self.auth_code = hash_object.hexdigest()

        messagebox.showinfo("Hash gerada", f"Código de Autenticação (hash): {self.auth_code}")
        self.signed_hashes.append(self.auth_code) 

    def browse_image(self):
        self.image_file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg *.gif *.bmp *.tiff")])

    def sign_file(self):
        if not self.file_path or not self.auth_code:
            messagebox.showerror("Error", "Please select a file and generate a hash first.")
            return
        
        if self.image_file_path:
            try:
                image = Image.open(self.image_file_path)
                image = image.convert("RGBA")
                image_reader = ImageReader(image)
            except:
                messagebox.showerror("Error", "Invalid image format or file.")
                return
        else:
            image_reader = None
        
        if self.file_path.endswith(".pdf"):
            self.sign_pdf(image_reader)
        elif self.file_path.endswith(".docx"):
            self.sign_docx(image_reader)
        else:
            messagebox.showerror("Error", "Invalid file format.")

    # ... (rest of the methods)

    def sign_pdf(self, image_reader):
        signer_name = self.signer_name_entry.get()
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        name_and_timestamp = f"Assinado por {signer_name}\nData: {timestamp}\nCódigo de Autenticação: {self.auth_code}"

        output_file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])

        pdf = PyPDF2.PdfReader(self.file_path)
        writer = PyPDF2.PdfWriter()

        for page_num in range(len(pdf.pages)):
            page = pdf.pages[page_num]
            
            # Create a PDF page
            new_page = PyPDF2.PageObject.create_blank_page(width=page.mediabox[2], height=page.mediabox[3])
            new_page.merge_page(page)

            # Draw text on the new page
            new_page.merge_page(self.create_text_object(name_and_timestamp, new_page))

            # Add background image if available
            if image_reader:
                new_page.merge_page(self.create_image_background(image_reader, new_page.mediabox[2], new_page.mediabox[3]))

            writer.add_page(new_page)
            self.progress_bar["value"] = (page_num + 1) * 100 / len(pdf.pages)
            self.root.update()

        with open(output_file_path, "wb") as output_file:
            writer.write(output_file)

        self.result_label.config(text="File signed and saved.")

    def sign_docx(self, image_reader):
        signer_name = self.signer_name_entry.get()
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        name_and_timestamp = f"Assinado por: {signer_name}\nData: {timestamp}\nCódigo de Autenticação: {self.auth_code}"

        doc = Document(self.file_path)
        
        if doc.paragraphs:
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.text += "\n" + name_and_timestamp

        if image_reader:
            image_stream = io.BytesIO()
            image = Image.open(self.image_file_path)
            image = image.convert("RGBA")
            image.save(image_stream, format="PNG")
            image_stream.seek(0)

            doc.add_picture(image_stream)

            self.progress_bar["value"] = 0  # Reset progress bar
            self.root.update()

        output_file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("DOCX Files", "*.docx")])
        doc.save(output_file_path)

        self.result_label.config(text="Documento assinado e salvo.")
        self.progress_bar["value"] = 100  # Set progress bar to 100%
        self.root.update()
        
    # Rest of the code remains the same
    def create_text_object(self, text, page):
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=(page.mediabox[2], page.mediabox[3]))
        text_lines = text.split('\n')
        
        font_size = 10  # Adjust the font size as needed
        
        # Calculate y_position for multi-line text
        line_height = 1.2 * font_size
        total_height = len(text_lines) * line_height
        y_position = 20 + total_height
        
        for line in text_lines:
            can.setFont("Helvetica", font_size)
            can.drawString(100, y_position, line)
            y_position -= line_height
        
        can.save()

        packet.seek(0)
        new_pdf = PyPDF2.PdfReader(packet)
        return new_pdf.pages[0]

    def create_image_background(self, image_reader, width, height):
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=(width, height))
        
        # Draw the image
        image_width, image_height = image_reader.getSize()
        image_position_x = (width - image_width) / 2
        image_position_y = (height - image_height) / 2
        can.drawImage(image_reader, image_position_x, image_position_y, width=image_width, height=image_height)
        
        can.save()

        packet.seek(0)
        new_pdf = PyPDF2.PdfReader(packet)
        return new_pdf.pages[0]

    def authenticate_hash(self):
        entered_hash = self.authenticator_entry.get()
        
        if entered_hash in self.signed_hashes:
            messagebox.showinfo("Autenticada com sucesso", "Hash é autêntica.")
        else:
            messagebox.showerror("Falha na autenticação", "Hash não é autêntica.")

app = tk.Tk()
PDFSignerApp(app)
